from typing import Generator, Iterable
import os
import docx
from docx import Document
import docx.document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from src.services.dados.arquivo import Arquivo


class Documento(Arquivo[docx.document.Document]):

    def __init__(self, nome_arquivo: str = None, texto: str = None, opcao_tratamento: int = None) -> None:
        super().__init__(nome_arquivo, texto)
        self.__fonte = 'Ubuntu'
        self.__cor = RGBColor(0, 0, 0)
        self.__alinhamento_justificado = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.__documento = self._abrir_arquivo()
        self.__opcao_tratamento = opcao_tratamento

    def ler_valores(self) -> Generator[Iterable[str], None, None]:
        return super().ler_valores()

    def _abrir_arquivo(self) -> docx.document.Document:
        documento = Document()
        return documento

    @property
    def opcao_tratamento(self):
        return self.__opcao_tratamento

    @opcao_tratamento.setter
    def opcao_tratamento(self, opcao_tratamento: int):
        self.__opcao_tratamento = opcao_tratamento

    def gravar_dados(self, **kwargs):
        if self.__opcao_tratamento == 1:
            linhas = self.texto.split("\n")
            for linha in linhas:
                if linha.startswith("##"):
                    titulo = self.__documento.add_heading(linha[3:], level=1)
                    for run in titulo.runs:
                        run.font.color.rgb = self.__cor
                        run.font.name = self.__fonte
                elif linha.startswith("###"):
                    subtitulo = self.__documento.add_heading(
                        linha[4:], level=2)
                    for run in subtitulo.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Ubuntu'
                elif linha.startswith("*"):

                    paragrafo = self.__documento.add_paragraph(
                        linha[2:], style='List Bullet')
                    paragrafo.alignment = self.__alinhamento_justificado
                    for run in paragrafo.runs:
                        run.font.name = self.__fonte

    def salvar_dados(self, nome_arquivo: str = None):
        caminho_arquivo_salvar = os.path.join(
            self._caminho_base, 'docs', nome_arquivo)
        self.__documento.save(caminho_arquivo_salvar)

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_value, traceback):
        self.__documento = self._abrir_arquivo()
